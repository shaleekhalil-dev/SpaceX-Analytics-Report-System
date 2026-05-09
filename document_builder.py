from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_word_report():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    title = doc.add_heading('SpaceX Flight Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "This report presents an analysis of rocket landing success factors "
        "based on a simulation model of Falcon 9 flight data."
    )

    doc.add_heading('2. Methodology', level=1)
    doc.add_paragraph(
        "A synthetic dataset of 100 flights was generated to study variables "
        "including Payload Mass, Orbit Type, and Landing Outcomes."
    )

    doc.add_heading('3. Visual Analysis', level=1)
    
    doc.add_heading('3.1 Success Rate by Orbit', level=2)
    if os.path.exists('figures/orbit_success_rate.png'):
        doc.add_picture('figures/orbit_success_rate.png', width=Inches(5.5))
        doc.add_paragraph("Figure 1: Landing success distribution across different orbits.")

    doc.add_page_break()

    doc.add_heading('3.2 Payload Mass Impact', level=2)
    if os.path.exists('figures/payload_vs_outcome.png'):
        doc.add_picture('figures/payload_vs_outcome.png', width=Inches(5.5))
        doc.add_paragraph("Figure 2: Statistical distribution of payload mass for success/failure outcomes.")

    doc.add_heading('4. Conclusion', level=1)
    doc.add_paragraph(
        "Initial analysis indicates that orbit type is a significant factor in landing success, "
        "while payload mass requires further predictive modeling."
    )

    output_path = 'outputs/SpaceX_Analysis_Report.docx'
    doc.save(output_path)
    print(f"Word document generated successfully at: {output_path}")

if __name__ == "__main__":
    create_word_report()